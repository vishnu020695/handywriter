"""
HandyWriter — Handwritten Image → Word/Excel + Simple PDF Editor
------------------------------------------------------------------
100% free, open-source Python tools only (Tesseract OCR, python-docx,
openpyxl, PyMuPDF). Runs as a local web app in your browser — works
on Windows, Mac, Linux, and Android/iPhone (via browser, same WiFi).

HOW TO RUN (see README.md for full details):
    pip install -r requirements.txt
    streamlit run app.py
"""

import io
import os
import re
import tempfile
import zipfile

import streamlit as st
from PIL import Image, ImageDraw, ImageFont
import pytesseract
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
import openpyxl

st.set_page_config(page_title="HandyWriter", page_icon="✍️", layout="wide")
st.title("✍️ HandyWriter")
st.caption("Convert handwritten/scanned images to Word or Excel, and edit PDFs — free & offline.")

tab1, tab2, tab3 = st.tabs(
    ["📝 Image → Word / Excel", "📄 PDF Editor", "📬 Bulk Mail Merge"]
)


# ---------------------------------------------------------------------------
# TAB 1: Image -> Word / Excel
# ---------------------------------------------------------------------------
with tab1:
    st.subheader("Convert a handwritten or scanned photo into an editable document")
    st.write(
        "Upload a clear photo of handwriting or printed text. The app reads the text "
        "automatically (OCR). **Always check the extracted text below before saving** — "
        "OCR is not 100% perfect, especially with cursive handwriting, so this lets you "
        "fix any mistakes first."
    )

    uploaded_img = st.file_uploader(
        "Upload image (JPG, PNG)", type=["jpg", "jpeg", "png"], key="img_upload"
    )

    if uploaded_img:
        image = Image.open(uploaded_img).convert("RGB")
        col1, col2 = st.columns(2)
        with col1:
            st.image(image, caption="Your uploaded image", use_container_width=True)

        with st.spinner("Reading text from image..."):
            # psm 6 = assume a uniform block of text (good default for notes/pages)
            extracted_text = pytesseract.image_to_string(image, config="--psm 6")

        with col2:
            st.write("**Extracted text (edit to fix any mistakes):**")
            edited_text = st.text_area(
                "Editable text", extracted_text, height=350, label_visibility="collapsed"
            )

        st.write("### Save as:")
        c1, c2 = st.columns(2)

        with c1:
            if st.button("💾 Create Word (.docx)", use_container_width=True):
                doc = Document()
                doc.add_heading("Converted from image", level=1)
                for line in edited_text.split("\n"):
                    doc.add_paragraph(line)
                buf = io.BytesIO()
                doc.save(buf)
                buf.seek(0)
                st.download_button(
                    "⬇️ Download Word file",
                    data=buf,
                    file_name="converted.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )

        with c2:
            if st.button("💾 Create Excel (.xlsx)", use_container_width=True):
                # Each non-empty line becomes a row; splits on 2+ spaces or tabs into columns
                wb = openpyxl.Workbook()
                ws = wb.active
                ws.title = "Converted"
                for line in edited_text.split("\n"):
                    if not line.strip():
                        continue
                    # naive column split: tabs, or 2+ spaces, treated as column breaks
                    import re
                    cells = re.split(r"\t|\s{2,}", line.strip())
                    ws.append(cells)
                buf = io.BytesIO()
                wb.save(buf)
                buf.seek(0)
                st.download_button(
                    "⬇️ Download Excel file",
                    data=buf,
                    file_name="converted.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                )

# ---------------------------------------------------------------------------
# TAB 2: PDF Editor
# ---------------------------------------------------------------------------
with tab2:
    st.subheader("Simple PDF editing (no misaligned pages, no corruption)")
    st.write(
        "Upload a PDF to merge, delete pages, rotate, add a watermark, or convert it into "
        "an editable Word document."
    )

    uploaded_pdf = st.file_uploader("Upload PDF", type=["pdf"], key="pdf_upload")

    if uploaded_pdf:
        pdf_bytes = uploaded_pdf.read()
        doc_pdf = fitz.open(stream=pdf_bytes, filetype="pdf")
        st.info(f"Loaded PDF with **{doc_pdf.page_count}** pages.")

        with st.expander(f"📖 Browse pages (see exact text on any of the {doc_pdf.page_count} pages)", expanded=(doc_pdf.page_count > 1)):
            browse_page_num = st.number_input(
                "Page to view", min_value=1, max_value=doc_pdf.page_count, value=1, key="browse_page_num"
            )
            _browse_page = doc_pdf[browse_page_num - 1]
            _browse_pix = _browse_page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
            st.image(_browse_pix.tobytes("png"), caption=f"Page {browse_page_num} of {doc_pdf.page_count}")
            with st.popover("📋 Show this page's raw text (to copy exact spelling/spacing)"):
                st.text(_browse_page.get_text())

        action = st.selectbox(
            "Choose an action",
            [
                "Find & Replace (exact text — safest, no misalignment)",
                "Cover & Stamp (fixes tricky fonts / broken text)",
                "Edit page text (direct, in-place)",
                "Edit images / logo (replace or delete)",
                "Delete pages",
                "Rotate pages",
                "Add text watermark",
                "Convert to editable Word (.docx)",
            ],
        )

        if action == "Cover & Stamp (fixes tricky fonts / broken text)":
            st.write(
                "Use this when a certificate/letter uses a fancy or unusual font that "
                "makes other tools misalign text (you'll see this if Find & Replace "
                "produced garbled or overlapping results). This draws a plain white box "
                "over the broken area and types new text at an exact position you choose "
                "— completely independent of the original file's font data, so it can't "
                "misalign."
            )
            stamp_page_num = st.number_input(
                "Page number", min_value=1, max_value=doc_pdf.page_count, value=1, key="stamp_page_num"
            )
            stamp_page_index = stamp_page_num - 1
            stamp_page = doc_pdf[stamp_page_index]
            stamp_page_rect = stamp_page.rect

            zoom = 1.5
            stamp_pix = stamp_page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
            stamp_img_w, stamp_img_h = stamp_pix.width, stamp_pix.height

            st.write("**Step 1: Set the cover box position** (in pixels, matching the preview image below)")
            cb1, cb2, cb3, cb4 = st.columns(4)
            with cb1:
                cov_x0 = st.number_input("Left (X0)", min_value=0, max_value=stamp_img_w, value=100, key="cov_x0")
            with cb2:
                cov_y0 = st.number_input("Top (Y0)", min_value=0, max_value=stamp_img_h, value=100, key="cov_y0")
            with cb3:
                cov_x1 = st.number_input("Right (X1)", min_value=0, max_value=stamp_img_w, value=400, key="cov_x1")
            with cb4:
                cov_y1 = st.number_input("Bottom (Y1)", min_value=0, max_value=stamp_img_h, value=130, key="cov_y1")

            # preview with red box overlay at current settings
            _prev_img = Image.open(io.BytesIO(stamp_pix.tobytes("png"))).convert("RGB")
            from PIL import ImageDraw as _ImgDraw2
            _draw2 = _ImgDraw2.Draw(_prev_img)
            _draw2.rectangle([cov_x0, cov_y0, cov_x1, cov_y1], outline="red", width=3)
            st.image(_prev_img, caption=f"Page {stamp_page_num} — red box shows where the cover + new text will go", width=600)

            st.write("**Step 2: New text and font settings**")
            stamp_text = st.text_area("Text to type in that box", key="stamp_text")
            sf1, sf2, sf3, sf4 = st.columns(4)
            with sf1:
                stamp_font_choice = st.selectbox(
                    "Font", ["Default (Helvetica)", "Times New Roman", "Courier (monospace)"], key="stamp_font_choice"
                )
            with sf2:
                stamp_bold = st.checkbox("Bold", key="stamp_bold")
            with sf3:
                stamp_size = st.number_input("Font size", min_value=6, max_value=200, value=13, key="stamp_size")
            with sf4:
                stamp_align = st.selectbox("Align", ["Left", "Center"], key="stamp_align")

            _stamp_font_map = {
                "Default (Helvetica)": {"regular": "helv", "bold": "hebo"},
                "Times New Roman": {"regular": "tiro", "bold": "tibo"},
                "Courier (monospace)": {"regular": "cour", "bold": "cobo"},
            }
            stamp_fontname = _stamp_font_map[stamp_font_choice]["bold" if stamp_bold else "regular"]
            stamp_align_code = 1 if stamp_align == "Center" else 0

            if st.button("Apply cover & stamp, then download"):
                if not stamp_text.strip():
                    st.warning("Enter the text to type into the box.")
                else:
                    # convert pixel coords (at preview zoom) back to PDF point coords
                    pdf_rect = fitz.Rect(cov_x0 / zoom, cov_y0 / zoom, cov_x1 / zoom, cov_y1 / zoom)
                    stamp_page.draw_rect(pdf_rect, color=None, fill=(1, 1, 1), fill_opacity=1)
                    rc = stamp_page.insert_textbox(
                        pdf_rect, stamp_text, fontsize=stamp_size, fontname=stamp_fontname, align=stamp_align_code
                    )
                    size = stamp_size
                    tries = 0
                    while rc < 0 and size > 6 and tries < 8:
                        size -= 1
                        rc = stamp_page.insert_textbox(
                            pdf_rect, stamp_text, fontsize=size, fontname=stamp_fontname, align=stamp_align_code
                        )
                        tries += 1
                    out_stamp = io.BytesIO(doc_pdf.tobytes())
                    st.success("Applied.")
                    st.download_button(
                        "⬇️ Download edited PDF",
                        data=out_stamp,
                        file_name="stamped_edited.pdf",
                        mime="application/pdf",
                    )

        elif action == "Find & Replace (exact text — safest, no misalignment)":
            st.write(
                "Type the exact text you want to change (e.g. `AJAY K`) and what to "
                "change it to. This only touches that exact text — everything else on "
                "the page, including text right next to it, stays frozen in place. "
                "This is the safest option when a name/value sits inside a longer "
                "sentence or line of underscores."
            )
            if "fr_count" not in st.session_state:
                st.session_state.fr_count = 2
            c_add, c_rem = st.columns(2)
            with c_add:
                if st.button("➕ Add another find & replace"):
                    st.session_state.fr_count += 1
                    st.rerun()
            with c_rem:
                if st.button("🗑️ Remove last") and st.session_state.fr_count > 1:
                    st.session_state.fr_count -= 1
                    st.rerun()

            fr_pairs = []
            for i in range(st.session_state.fr_count):
                cfind, creplace = st.columns(2)
                with cfind:
                    find_val = st.text_input(f"Find (exact text) #{i+1}", key=f"fr_find_{i}")
                with creplace:
                    replace_val = st.text_input(f"Replace with #{i+1}", key=f"fr_replace_{i}")
                if find_val.strip():
                    fr_pairs.append((find_val, replace_val))

            st.write("**Font settings for the replacement text:**")
            cf1, cf2, cf3 = st.columns(3)
            with cf1:
                font_choice = st.selectbox(
                    "Font", ["Default (Helvetica)", "Times New Roman", "Courier (monospace)"],
                    key="fr_font_choice",
                )
            with cf2:
                fr_bold = st.checkbox("Bold", key="fr_bold")
            with cf3:
                fr_size_override = st.number_input(
                    "Font size (0 = auto-match original)", min_value=0, max_value=200, value=0, key="fr_size"
                )

            _font_map = {
                "Default (Helvetica)": {"regular": "helv", "bold": "hebo"},
                "Times New Roman": {"regular": "tiro", "bold": "tibo"},
                "Courier (monospace)": {"regular": "cour", "bold": "cobo"},
            }
            fr_fontname = _font_map[font_choice]["bold" if fr_bold else "regular"]

            if st.button("🔍 Preview matches (see which pages, before replacing)"):
                if not fr_pairs:
                    st.warning("Enter at least one 'Find' value.")
                else:
                    any_found = False
                    for find_val, _ in fr_pairs:
                        pages_with_match = []
                        for pnum in range(doc_pdf.page_count):
                            rects = doc_pdf[pnum].search_for(find_val)
                            if rects:
                                pages_with_match.append((pnum + 1, len(rects)))
                        if pages_with_match:
                            any_found = True
                            st.write(f"**\"{find_val}\"** found on: " + ", ".join(f"page {p} ({n}x)" for p, n in pages_with_match))
                            first_page_num, _ = pages_with_match[0]
                            _prev_page = doc_pdf[first_page_num - 1]
                            _prev_rects = _prev_page.search_for(find_val)
                            zoom = 1.5
                            _pix = _prev_page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
                            _img = Image.open(io.BytesIO(_pix.tobytes("png"))).convert("RGB")
                            from PIL import ImageDraw as _ImageDraw
                            _draw = _ImageDraw.Draw(_img)
                            for r in _prev_rects:
                                _draw.rectangle([r.x0*zoom, r.y0*zoom, r.x1*zoom, r.y1*zoom], outline="red", width=3)
                            st.image(_img, caption=f"Highlighted on page {first_page_num}", width=500)
                        else:
                            st.warning(f"**\"{find_val}\"** was not found anywhere in the document. Check spelling/spacing.")
                    if not any_found:
                        st.info("No matches found for any entry — nothing will change if you apply now.")

            def _flags_to_base14(flags, serif_hint):
                """Guess the closest Base-14 font from PyMuPDF span flags."""
                italic = bool(flags & 2)
                bold = bool(flags & 16)
                mono = bool(flags & 8)
                serif = serif_hint or bool(flags & 4)
                if mono:
                    if bold and italic: return "cobi"
                    if bold: return "cobo"
                    if italic: return "coit"
                    return "cour"
                if serif:
                    if bold and italic: return "tibi"
                    if bold: return "tibo"
                    if italic: return "tiit"
                    return "tiro"
                if bold and italic: return "hebi"
                if bold: return "hebo"
                if italic: return "heit"
                return "helv"

            def _find_line_for_rect(page, rect):
                """Return the full text line dict that contains rect."""
                text_dict = page.get_text("dict")
                cy = (rect.y0 + rect.y1) / 2
                best = None
                for block in text_dict.get("blocks", []):
                    for line in block.get("lines", []):
                        bbox = line["bbox"]
                        if bbox[1] - 1 <= cy <= bbox[3] + 1:
                            best = line
                            break
                    if best:
                        break
                return best

            font_buffer_cache = {}

            def _embedded_font_alias(page, span_font_name, want_bold=False):
                """Try to pull the ACTUAL embedded font used by the original text,
                so the replacement matches the certificate's real typeface instead
                of falling back to a generic Base-14 font. If want_bold is True,
                only accept a font that is actually bold (either a genuine bold
                sibling like 'Arial-Bold', or the original font itself if IT is
                already the bold weight) — never silently hand back a regular
                weight when the user asked for bold."""
                try:
                    candidates = list(page.get_fonts(full=True))
                    family_root = (span_font_name or "").split("-")[0].split(",")[0].lower()

                    if want_bold:
                        # 1) look for a genuine bold sibling of the same family
                        for xref, ext, subtype, basefont, refname, encoding in candidates:
                            bf_lower = basefont.lower()
                            if family_root and family_root in bf_lower and "bold" in bf_lower:
                                fontname, fontext, fonttype, fontbuffer = doc_pdf.extract_font(xref)
                                if fontbuffer:
                                    alias = f"embedded_{xref}"
                                    page.insert_font(fontname=alias, fontbuffer=fontbuffer)
                                    font_buffer_cache[alias] = fontbuffer
                                    return alias
                        # 2) the original font itself is already bold-named
                        if span_font_name and "bold" in span_font_name.lower():
                            for xref, ext, subtype, basefont, refname, encoding in candidates:
                                if basefont == span_font_name:
                                    fontname, fontext, fonttype, fontbuffer = doc_pdf.extract_font(xref)
                                    if fontbuffer:
                                        alias = f"embedded_{xref}"
                                        page.insert_font(fontname=alias, fontbuffer=fontbuffer)
                                        font_buffer_cache[alias] = fontbuffer
                                        return alias
                        # no real bold variant embedded in this PDF — tell the caller
                        # so it can fall back to a synthetic bold instead of silently
                        # rendering regular weight
                        return None

                    # regular weight: reuse the exact original font
                    for xref, ext, subtype, basefont, refname, encoding in candidates:
                        if basefont == span_font_name or (span_font_name and span_font_name in basefont):
                            fontname, fontext, fonttype, fontbuffer = doc_pdf.extract_font(xref)
                            if fontbuffer:
                                alias = f"embedded_{xref}"
                                page.insert_font(fontname=alias, fontbuffer=fontbuffer)
                                font_buffer_cache[alias] = fontbuffer
                                return alias
                except Exception:
                    pass
                return None

            def _measure_width(text, fontname, fontsize):
                """Measure text width using the ACTUAL font it will be drawn
                in, not a generic Helvetica guess. Using the wrong font's
                metrics is what caused the width estimate to be badly off —
                too small (text got shrunk to fit) or, when over-compensated
                with a flat safety margin, too large (the box spilled into
                neighboring real text like a dash or the next word)."""
                try:
                    if fontname in font_buffer_cache:
                        return fitz.Font(fontbuffer=font_buffer_cache[fontname]).text_length(text, fontsize=fontsize)
                    return fitz.Font(fontname=fontname).text_length(text, fontsize=fontsize)
                except Exception:
                    return fitz.get_text_length(text, fontname="helv", fontsize=fontsize)

            def _is_filler_span(text):
                """Detect blank-line filler such as long underscore runs used
                for 'fill in the blank' fields (e.g. the name/department lines
                on a certificate)."""
                stripped = text.replace(" ", "")
                if not stripped:
                    return True
                filler_chars = set("_.-")
                return len(stripped) >= 3 and all(ch in filler_chars for ch in stripped)

            def _adjacent_fillers(line, rect):
                """Return the immediately adjacent filler spans (left/right) of
                the span containing rect, as (text, bbox, font, size) tuples —
                so callers can re-queue them for redraw whenever a wider
                replacement's box might encroach on that blank-line ink."""
                if not line or not line.get("spans"):
                    return []
                spans = line["spans"]
                match_idx = None
                for i, s in enumerate(spans):
                    if fitz.Rect(s["bbox"]).intersects(rect):
                        match_idx = i
                        break
                if match_idx is None:
                    return []
                out = []
                if match_idx - 1 >= 0 and _is_filler_span(spans[match_idx - 1]["text"]):
                    s = spans[match_idx - 1]
                    out.append((s["text"], fitz.Rect(s["bbox"]), s.get("font"), s["size"]))
                if match_idx + 1 < len(spans) and _is_filler_span(spans[match_idx + 1]["text"]):
                    s = spans[match_idx + 1]
                    out.append((s["text"], fitz.Rect(s["bbox"]), s.get("font"), s["size"]))
                return out

            if st.button("Apply Find & Replace to entire PDF"):
                if not fr_pairs:
                    st.warning("Enter at least one 'Find' value.")
                else:
                    total_replacements = 0
                    for page in doc_pdf:
                        page_rect = page.rect
                        annots_queued = 0
                        # Queue EVERY match for EVERY find/replace pair FIRST,
                        # all measured against the original, untouched page —
                        # then apply them all in one atomic pass at the end.
                        # This is the key fix: nothing is redrawn until every
                        # match has already been located, so an earlier
                        # replacement can never shift the page and invalidate
                        # a later match's coordinates (which is what caused
                        # the repeated "202 2027 202 2027" corruption before).
                        for find_val, replace_val in fr_pairs:
                            rects = page.search_for(find_val)
                            for rect in rects:
                                line = _find_line_for_rect(page, rect)
                                fontsize = 11
                                span_font_name = None
                                span_flags = 0
                                if line and line.get("spans"):
                                    for s in line["spans"]:
                                        if fitz.Rect(s["bbox"]).intersects(rect):
                                            fontsize = s["size"]
                                            span_font_name = s.get("font")
                                            span_flags = s.get("flags", 0)
                                            break

                                # Sanity-check the detected font size against the
                                # match's own rectangle height (a hard geometric
                                # fact from search_for, always reliable). If the
                                # span lookup returned something implausibly
                                # small — e.g. it grabbed a neighboring span with
                                # different metrics — this is what made replaced
                                # numbers render tiny and float like superscript.
                                height_based_size = rect.height / 1.15
                                if fontsize < height_based_size * 0.7:
                                    fontsize = height_based_size

                                if fr_size_override > 0:
                                    fontsize = fr_size_override

                                use_fontname = fr_fontname
                                if font_choice == "Default (Helvetica)" and span_font_name:
                                    embedded_alias = _embedded_font_alias(page, span_font_name, want_bold=fr_bold)
                                    if embedded_alias:
                                        use_fontname = embedded_alias
                                    else:
                                        effective_flags = span_flags | 16 if fr_bold else span_flags
                                        use_fontname = _flags_to_base14(effective_flags, serif_hint=True)

                                # Measure the replacement using its ACTUAL font
                                # (embedded font or the real Base-14 we picked),
                                # not a generic Helvetica guess — that mismatch
                                # is what previously caused either shrinking
                                # (underestimate) or spilling into neighboring
                                # text like a dash or the next word (overshoot).
                                # No flat safety margin is added anymore: since
                                # the measurement is now accurate, any extra
                                # padding just shows up as a visible blank gap
                                # before whatever character comes right after
                                # (e.g. "2066  -" instead of "2066 -") — which
                                # is exactly the kind of seam that makes an
                                # edit look like an edit. Only widen the box
                                # when the new text genuinely needs more room.
                                est_w = _measure_width(replace_val, use_fontname, fontsize)
                                extra_w = max(0, est_w - rect.width)
                                # Vertical padding is asymmetric on purpose:
                                # generous ABOVE the text (room for ascenders),
                                # almost none BELOW it — a blank "fill in the
                                # blank" line is very often a drawn rule
                                # graphic just under the baseline, not
                                # underscore text, so the safest way to avoid
                                # ever painting over it is to keep the box's
                                # bottom edge close to the text itself.
                                vpad_top = fontsize * 0.3
                                vpad_bottom = fontsize * 0.05
                                box_x1 = min(rect.x1 + extra_w, page_rect.width - 5)
                                box = fitz.Rect(
                                    rect.x0, rect.y0 - vpad_top,
                                    box_x1,
                                    rect.y1 + vpad_bottom,
                                )

                                # If this widened box now reaches into an
                                # adjacent blank-line filler (underscores),
                                # queue that filler for redraw too — unchanged
                                # text, same spot — so the blank line stays
                                # continuous instead of losing ink where the
                                # box overlapped it.
                                for f_text, f_bbox, f_font, f_size in _adjacent_fillers(line, rect):
                                    if box.intersects(f_bbox):
                                        f_alias = _embedded_font_alias(page, f_font, want_bold=False) or "tiro"
                                        page.add_redact_annot(
                                            f_bbox,
                                            text=f_text,
                                            fontname=f_alias,
                                            fontsize=f_size,
                                            align=fitz.TEXT_ALIGN_LEFT,
                                            fill=(1, 1, 1),
                                            text_color=(0, 0, 0),
                                        )
                                        annots_queued += 1

                                page.add_redact_annot(
                                    box,
                                    text=replace_val,
                                    fontname=use_fontname,
                                    fontsize=fontsize,
                                    align=fitz.TEXT_ALIGN_LEFT,
                                    fill=(1, 1, 1),
                                    text_color=(0, 0, 0),
                                )
                                annots_queued += 1
                                total_replacements += 1

                        if annots_queued:
                            page.apply_redactions()

                    if total_replacements == 0:
                        st.warning("None of the 'Find' text was found in this PDF. Check spelling/spacing matches exactly.")
                    else:
                        out = io.BytesIO(doc_pdf.tobytes())
                        st.success(f"Made {total_replacements} replacement(s) across the document.")
                        st.download_button(
                            "⬇️ Download edited PDF",
                            data=out,
                            file_name="find_replace_edited.pdf",
                            mime="application/pdf",
                        )

        elif action == "Edit page text (direct, in-place)":
            st.write(
                "Works best on PDFs that already have real text (not scanned photos). "
                "Each numbered box below matches the same numbered box drawn on the page "
                "preview image, so you can see exactly where your edit will land. "
                "Edit any box, then apply — text is placed in the same spot and size, "
                "so nothing shifts or misaligns."
            )

            st.markdown("##### 🔍 Find it first — search every page")
            st.caption(
                "On a multi-page document it's much faster to search across all pages "
                "than to open each one and hunt through its boxes."
            )
            global_query = st.text_input(
                "Search across ALL pages for text you want to edit", key="el_global_query"
            )
            if global_query.strip():
                q = global_query.strip().lower()
                results = []  # (page_index, snippet)
                for pi in range(doc_pdf.page_count):
                    pg_td = doc_pdf[pi].get_text("dict")
                    for db in pg_td["blocks"]:
                        for ln in db.get("lines", []):
                            txt = "".join(s["text"] for s in ln["spans"]).strip()
                            if txt and q in txt.lower():
                                results.append((pi, txt))
                if not results:
                    st.warning("That text wasn't found on any page.")
                else:
                    pages_hit = sorted(set(r[0] + 1 for r in results))
                    st.write(f"Found on page(s): **{', '.join(str(p) for p in pages_hit)}** "
                             f"({len(results)} matching line(s) total). Click one to jump straight there:")
                    for result_idx, (pi, snippet) in enumerate(results[:40]):
                        label = f"Page {pi + 1}: {snippet[:80]}"
                        if st.button(label, key=f"jump_{result_idx}_{pi}"):
                            st.session_state["el_page_num"] = pi + 1
                            st.session_state[f"line_search_{pi}"] = global_query
                            st.rerun()
                    if len(results) > 40:
                        st.caption(f"...and {len(results) - 40} more matches not shown here.")
            st.markdown("---")

            if "el_page_num" not in st.session_state:
                st.session_state["el_page_num"] = 1
            page_num = st.number_input(
                "Page number to edit", min_value=1, max_value=doc_pdf.page_count, key="el_page_num"
            )
            page_index = page_num - 1
            page = doc_pdf[page_index]

            # Extract editable text at the LINE level (not paragraph blocks, not
            # individual words). This is the level that correctly keeps separate
            # fields on the same visual line (e.g. a name and a department sitting
            # side by side on a certificate) as separate editable boxes, while still
            # keeping normal sentences/paragraphs as single boxes.
            text_dict = page.get_text("dict")
            lines_data = []  # list of (bbox, text, fontsize, font_name, flags, is_mixed_style)
            for db in text_dict["blocks"]:
                for line in db.get("lines", []):
                    line_text = "".join(s["text"] for s in line["spans"])
                    if line_text.strip():
                        spans = line["spans"]
                        if spans:
                            # Use the DOMINANT span (most characters), not just
                            # the first one. A sentence like "has successfully
                            # completed **Enhancement Program**" starts with
                            # regular text but the first span alone isn't
                            # representative — picking font/bold/size from
                            # whichever span happens to come first was a
                            # coin-flip that could make an entire mostly-plain
                            # sentence render bold if it happened to start on
                            # a short bold fragment (or vice versa).
                            dominant = max(spans, key=lambda s: len(s.get("text", "")))
                            fontsize = dominant["size"]
                            font_name = dominant.get("font")
                            flags = dominant.get("flags", 0)

                            # A single box can only ever be drawn in ONE font.
                            # If this line genuinely mixes styles (e.g. plain
                            # text with a bold heading embedded in the middle)
                            # with more than one style contributing a real
                            # share of the text, no amount of "pick the best
                            # span" logic can render it correctly — the
                            # non-dominant portion will always end up wrong.
                            # Detect that case so we can warn instead of
                            # silently flattening it.
                            style_chars = {}
                            total_chars = 0
                            for s in spans:
                                key = (s.get("font"), bool(s.get("flags", 0) & 16))
                                n = len(s.get("text", ""))
                                style_chars[key] = style_chars.get(key, 0) + n
                                total_chars += n
                            is_mixed_style = total_chars > 0 and sum(
                                1 for v in style_chars.values() if v / total_chars >= 0.15
                            ) > 1
                        else:
                            fontsize, font_name, flags, is_mixed_style = 11, None, 0, False
                        lines_data.append((line["bbox"], line_text, fontsize, font_name, flags, is_mixed_style))

            search_query = st.text_input(
                "🔍 Type part of the text you want to change (leave empty to see every box)",
                key=f"line_search_{page_index}",
            )
            if search_query.strip():
                matching_idx = {
                    i for i, (_, text, _, _, _, _) in enumerate(lines_data)
                    if search_query.strip().lower() in text.lower()
                }
                if not matching_idx:
                    st.warning("No text on this page matches that — showing every box instead.")
                    matching_idx = set(range(len(lines_data)))
            else:
                matching_idx = set(range(len(lines_data)))

            # Draw numbered boxes on the preview so boxes map visually to the page.
            # Matches (or everything, if no search yet) are outlined in red and
            # numbered; anything filtered out by the search is a faint gray so the
            # page doesn't turn into unreadable clutter but you can still see where
            # everything sits.
            zoom = 1.3
            pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
            from PIL import Image as PILImage, ImageDraw
            preview_img = PILImage.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
            draw = ImageDraw.Draw(preview_img)
            for i, (bbox, _, _, _, _, _) in enumerate(lines_data):
                x0, y0, x1, y1 = [v * zoom for v in bbox]
                if i in matching_idx:
                    draw.rectangle([x0, y0, x1, y1], outline="red", width=2)
                    draw.text((x0, max(0, y0 - 14)), f"#{i + 1}", fill="red")
                else:
                    draw.rectangle([x0, y0, x1, y1], outline=(200, 200, 200), width=1)
            st.image(preview_img, caption=f"Page {page_num} — boxes numbered to match the list below", width=500)

            if not lines_data:
                st.warning(
                    "No selectable text found on this page — it's likely a scanned image. "
                    "Use the 'Image → Word/Excel' tab instead, or the PDF-to-Word conversion below."
                )
            else:
                shown_count = len(matching_idx)
                if search_query.strip():
                    st.write(f"**{shown_count}** of {len(lines_data)} boxes match — only these are shown below:")
                else:
                    st.write(f"**{len(lines_data)}** editable text lines found on this page:")
                edited_values = []
                for i, (bbox, line_text, fontsize, font_name, flags, is_mixed_style) in enumerate(lines_data):
                    original_text = line_text.rstrip("\n")
                    if i not in matching_idx:
                        # keep the value unchanged for boxes hidden by the filter,
                        # but don't render a text_area for them so the list stays short
                        edited_values.append(original_text)
                        continue
                    if is_mixed_style:
                        st.warning(
                            f"⚠️ Box #{i + 1} mixes bold and regular text in the same line. "
                            "Editing it here will flatten the WHOLE box to one style. "
                            "If you only need to change a small part (like a date or a "
                            "single word) and want the bold/regular parts to stay correct, "
                            "use **Find & Replace** instead for just that part."
                        )
                    val = st.text_area(
                        f"Box #{i + 1}", original_text, height=68, key=f"line_{page_index}_{i}"
                    )
                    edited_values.append(val)

                st.write("**Font for edited boxes:**")
                lf1, lf2, lf3 = st.columns(3)
                with lf1:
                    line_font_choice = st.selectbox(
                        "Font", ["Auto-match original", "Force Helvetica", "Force Times New Roman", "Force Courier (monospace)"],
                        key=f"line_font_choice_{page_index}",
                        help="Auto-match tries to reuse each box's own original font "
                             "(so a bold heading stays bold Times, a plain line stays plain, etc). "
                             "Only pick a forced font if you specifically want to override that.",
                    )
                with lf2:
                    line_bold = st.checkbox("Force Bold", key=f"line_bold_{page_index}")
                with lf3:
                    line_size_override = st.number_input(
                        "Font size (0 = auto-match original)", min_value=0, max_value=200, value=0,
                        key=f"line_size_{page_index}",
                    )
                _line_font_map = {
                    "Force Helvetica": {"regular": "helv", "bold": "hebo"},
                    "Force Times New Roman": {"regular": "tiro", "bold": "tibo"},
                    "Force Courier (monospace)": {"regular": "cour", "bold": "cobo"},
                }

                def _line_flags_to_base14(flags, serif_hint, force_bold):
                    italic = bool(flags & 2)
                    bold = bool(flags & 16) or force_bold
                    mono = bool(flags & 8)
                    serif = serif_hint or bool(flags & 4)
                    if mono:
                        return "cobi" if (bold and italic) else "cobo" if bold else "coit" if italic else "cour"
                    if serif:
                        return "tibi" if (bold and italic) else "tibo" if bold else "tiit" if italic else "tiro"
                    return "hebi" if (bold and italic) else "hebo" if bold else "heit" if italic else "helv"

                _line_font_cache = {}

                def _line_embedded_font_alias(page, span_font_name):
                    """Reuse the box's ACTUAL original embedded font, so text
                    that's already the right typeface (e.g. bold Times for a
                    heading) doesn't get flattened into generic Helvetica —
                    that mismatch is what made an edited sentence look like a
                    pasted-in block before."""
                    if not span_font_name:
                        return None
                    if span_font_name in _line_font_cache:
                        return _line_font_cache[span_font_name]
                    try:
                        for xref, ext, subtype, basefont, refname, encoding in page.get_fonts(full=True):
                            if basefont == span_font_name or span_font_name in basefont:
                                fontname, fontext, fonttype, fontbuffer = doc_pdf.extract_font(xref)
                                if fontbuffer:
                                    alias = f"lineembed_{xref}"
                                    page.insert_font(fontname=alias, fontbuffer=fontbuffer)
                                    _line_font_cache[span_font_name] = alias
                                    return alias
                    except Exception:
                        pass
                    _line_font_cache[span_font_name] = None
                    return None

                def _split_filler_ends(text):
                    """Split off any leading/trailing blank-line filler
                    (underscore runs) from a line's text, returning
                    (leading_filler, core, trailing_filler). This lets us put
                    the blank-line ends back automatically even if the user's
                    edit didn't retype them — which is what silently erased
                    the underline under a name before: the box for a line is
                    its ENTIRE row including any underscores on either side,
                    so typing a plain name over that row drops them unless we
                    explicitly restore them here."""
                    core = text
                    leading = trailing = ""
                    m1 = re.match(r'^([_\s]{3,})', core)
                    if m1 and set(m1.group(1).replace(" ", "")) == {"_"}:
                        leading = m1.group(1)
                        core = core[len(leading):]
                    m2 = re.search(r'([_\s]{3,})$', core)
                    if m2 and set(m2.group(1).replace(" ", "")) == {"_"}:
                        trailing = m2.group(1)
                        core = core[:len(core) - len(trailing)]
                    return leading, core, trailing

                if st.button("Apply edits to this page and download"):
                    changed_any = False
                    edits = []  # (bbox, final_text, fontsize, fontname, is_blank_field)
                    for (bbox, line_text, fontsize, font_name, flags, is_mixed_style), new_val in zip(lines_data, edited_values):
                        original_text = line_text.rstrip("\n")
                        if new_val != original_text:
                            changed_any = True
                            use_size = line_size_override if line_size_override > 0 else fontsize

                            # Preserve the original blank-line underscores
                            # automatically unless the user's own edit already
                            # supplied its own leading/trailing filler.
                            orig_leading, _, orig_trailing = _split_filler_ends(original_text)
                            new_leading, new_core, new_trailing = _split_filler_ends(new_val)
                            final_leading = new_leading if new_leading else orig_leading
                            final_trailing = new_trailing if new_trailing else orig_trailing
                            final_val = final_leading + new_core + final_trailing
                            is_blank_field = bool(orig_leading) or bool(orig_trailing)

                            edits.append((fitz.Rect(bbox), final_val, use_size, font_name, flags, is_blank_field))

                    if not changed_any:
                        st.info("No changes were made.")
                    else:
                        page_rect = page.rect
                        underline_specs = []  # (x0, x1, y) to draw AFTER redactions apply
                        for bbox, final_val, fontsize, font_name, flags, is_blank_field in edits:
                            # Pick the font for THIS box: auto-match the box's
                            # own original font by default (reusing the actual
                            # embedded font when possible, else a flags-based
                            # Base-14 guess), or use the user's forced choice
                            # if they explicitly picked one.
                            if line_font_choice == "Auto-match original":
                                alias = _line_embedded_font_alias(page, font_name)
                                use_fontname = alias or _line_flags_to_base14(flags, serif_hint=True, force_bold=line_bold)
                            else:
                                use_fontname = _line_font_map[line_font_choice]["bold" if line_bold else "regular"]

                            # Measure with the ACTUAL font being used (accurate
                            # for Base-14 fonts) instead of assuming a fixed
                            # 300pt-wide box — that old approach either clipped
                            # long text or left a big empty gap for short text.
                            # Note: final_val already has any preserved filler
                            # (underscores) folded in by _split_filler_ends, so
                            # sizing the box to est_w alone is correct — do NOT
                            # also force a minimum of the original bbox.width,
                            # that was what left a big empty white rectangle
                            # whenever the new text was shorter than the old
                            # full line (and could even bleed into neighboring
                            # text/underscores that were never meant to be
                            # touched by this edit).
                            try:
                                est_w = fitz.get_text_length(final_val, fontname=use_fontname, fontsize=fontsize)
                            except Exception:
                                est_w = fitz.get_text_length(final_val, fontname="helv", fontsize=fontsize)
                            # Padding is asymmetric on purpose: generous ABOVE
                            # the text (room for ascenders/accents), but almost
                            # none BELOW it. A blank "fill in the name" line on
                            # a certificate is very often a drawn rule/line
                            # graphic sitting just under the baseline, not
                            # underscore characters — no amount of text-based
                            # filler detection can protect a line like that,
                            # since it isn't text at all. The only real
                            # protection is to make sure our white redaction
                            # fill never reaches down far enough to paint over
                            # it in the first place.
                            available_w = page_rect.width - 10 - bbox.x0
                            extra_h = 0
                            if est_w > available_w > 0:
                                # text will need to wrap onto more than one
                                # line within the box — give it the room, or
                                # PyMuPDF shrinks/clips it to force a fit.
                                wrap_lines = -(-int(est_w) // int(available_w))  # ceil
                                extra_h = (wrap_lines - 1) * fontsize * 1.3
                            box = fitz.Rect(
                                bbox.x0, bbox.y0 - fontsize * 0.3,
                                min(bbox.x0 + est_w + 4, page_rect.width - 10),
                                bbox.y1 + fontsize * 0.05 + extra_h,
                            )
                            page.add_redact_annot(
                                box,
                                text=final_val,
                                fontname=use_fontname,
                                fontsize=fontsize,
                                align=fitz.TEXT_ALIGN_LEFT,
                                fill=(1, 1, 1),
                                text_color=(0, 0, 0),
                            )
                            if is_blank_field:
                                # Draw a real underline back in ourselves,
                                # spanning the field's original full width,
                                # instead of hoping we correctly preserved
                                # whatever the original blank line was made
                                # of. This works regardless of whether that
                                # original line was underscore text or (more
                                # likely, since text-based preservation kept
                                # failing) a drawn rule graphic — we're not
                                # trying to protect it anymore, just putting a
                                # visible line back afterwards, unconditionally.
                                underline_specs.append((bbox.x0, bbox.x1, bbox.y1 + fontsize * 0.12))

                        page.apply_redactions()
                        for ux0, ux1, uy in underline_specs:
                            page.draw_line(fitz.Point(ux0, uy), fitz.Point(ux1, uy), color=(0, 0, 0), width=0.6)
                        out = io.BytesIO(doc_pdf.tobytes())
                        st.success("Edits applied.")
                        st.download_button(
                            "⬇️ Download edited PDF",
                            data=out,
                            file_name="edited_text.pdf",
                            mime="application/pdf",
                        )

        elif action == "Edit images / logo (replace or delete)":
            st.write(
                "Replace a logo/image with your own, or delete it entirely, without "
                "shifting any of the surrounding text."
            )
            page_num_img = st.number_input(
                "Page number", min_value=1, max_value=doc_pdf.page_count, value=1, key="img_page_num"
            )
            page_index_img = page_num_img - 1
            page_img = doc_pdf[page_index_img]

            images_found = page_img.get_images(full=True)
            if not images_found:
                st.warning("No images found on this page.")
            else:
                st.write(f"**{len(images_found)}** image(s) found on this page:")
                for idx, img_info in enumerate(images_found):
                    xref = img_info[0]
                    rects = page_img.get_image_rects(xref)
                    if not rects:
                        continue
                    rect = rects[0]

                    # show a small preview crop of this image area
                    zoom = 2
                    pix = page_img.get_pixmap(clip=rect, matrix=fitz.Matrix(zoom, zoom))
                    st.image(pix.tobytes("png"), caption=f"Image #{idx + 1} (page {page_num_img})", width=200)

                    col_a, col_b = st.columns(2)
                    with col_a:
                        replacement = st.file_uploader(
                            f"Replace image #{idx + 1} with:",
                            type=["png", "jpg", "jpeg"],
                            key=f"img_replace_{page_index_img}_{idx}",
                        )
                        if replacement and st.button(f"Apply replacement to image #{idx + 1}", key=f"apply_replace_{page_index_img}_{idx}"):
                            new_bytes = replacement.read()
                            page_img.add_redact_annot(rect, fill=(1, 1, 1))
                            page_img.apply_redactions()
                            page_img.insert_image(rect, stream=new_bytes)
                            out_img = io.BytesIO(doc_pdf.tobytes())
                            st.success(f"Image #{idx + 1} replaced.")
                            st.download_button(
                                "⬇️ Download edited PDF",
                                data=out_img,
                                file_name="edited_image.pdf",
                                mime="application/pdf",
                                key=f"dl_replace_{page_index_img}_{idx}",
                            )
                    with col_b:
                        if st.button(f"🗑️ Delete image #{idx + 1}", key=f"delete_{page_index_img}_{idx}"):
                            page_img.add_redact_annot(rect, fill=(1, 1, 1))
                            page_img.apply_redactions()
                            out_img = io.BytesIO(doc_pdf.tobytes())
                            st.success(f"Image #{idx + 1} deleted.")
                            st.download_button(
                                "⬇️ Download edited PDF",
                                data=out_img,
                                file_name="edited_image.pdf",
                                mime="application/pdf",
                                key=f"dl_delete_{page_index_img}_{idx}",
                            )

        elif action == "Delete pages":
            pages_to_delete = st.text_input(
                "Page numbers to delete (comma-separated, e.g. 1,3,5)"
            )
            if st.button("Apply and download"):
                try:
                    nums = [int(x.strip()) - 1 for x in pages_to_delete.split(",") if x.strip()]
                    doc_pdf.delete_pages(nums)
                    out = io.BytesIO(doc_pdf.tobytes())
                    st.download_button(
                        "⬇️ Download edited PDF", data=out, file_name="edited.pdf", mime="application/pdf"
                    )
                except Exception as e:
                    st.error(f"Check your page numbers. Details: {e}")

        elif action == "Rotate pages":
            angle = st.selectbox("Rotate by", [90, 180, 270])
            if st.button("Apply and download"):
                for page in doc_pdf:
                    page.set_rotation(angle)
                out = io.BytesIO(doc_pdf.tobytes())
                st.download_button(
                    "⬇️ Download rotated PDF", data=out, file_name="rotated.pdf", mime="application/pdf"
                )

        elif action == "Add text watermark":
            wm_text = st.text_input("Watermark text", "CONFIDENTIAL")
            if st.button("Apply and download"):
                for page in doc_pdf:
                    rect = page.rect
                    center = fitz.Point(rect.width / 2, rect.height / 2)
                    # rotate diagonally around the page center using a morph matrix
                    morph = (center, fitz.Matrix(45))
                    page.insert_text(
                        (rect.width / 4, rect.height / 2),
                        wm_text,
                        fontsize=40,
                        color=(0.7, 0.7, 0.7),
                        overlay=True,
                        morph=morph,
                    )
                out = io.BytesIO(doc_pdf.tobytes())
                st.download_button(
                    "⬇️ Download watermarked PDF", data=out, file_name="watermarked.pdf", mime="application/pdf"
                )

        elif action == "Convert to editable Word (.docx)":
            if st.button("Convert and download"):
                from pdf2docx import Converter

                with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_pdf:
                    tmp_pdf.write(pdf_bytes)
                    tmp_pdf_path = tmp_pdf.name
                tmp_docx_path = tmp_pdf_path.replace(".pdf", ".docx")

                with st.spinner("Converting (this keeps original layout/alignment)..."):
                    cv = Converter(tmp_pdf_path)
                    cv.convert(tmp_docx_path)
                    cv.close()

                with open(tmp_docx_path, "rb") as f:
                    st.download_button(
                        "⬇️ Download Word file",
                        data=f.read(),
                        file_name="converted_from_pdf.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                os.remove(tmp_pdf_path)
                os.remove(tmp_docx_path)

# ---------------------------------------------------------------------------
# TAB 3: Bulk Mail Merge (Excel + PDF template -> many personalized PDFs)
# ---------------------------------------------------------------------------
with tab3:
    st.subheader("Generate hundreds/thousands of personalized letters at once")

    merge_mode = st.radio(
        "Choose a method",
        [
            "Token-based ({{NAME}} placeholders in a PDF with real text)",
            "Coordinate-based (works with ANY template — PDF or image, no placeholders needed)",
        ],
    )

    if merge_mode.startswith("Token-based"):
        st.write(
            "**Step 1:** Prepare your template PDF with placeholder tokens where "
            "personal info should go — for example `{{NAME}}` and `{{DEPARTMENT}}`. "
            "Tip: you can create these tokens using the **PDF Editor** tab above — "
            "edit any real name/department in your existing offer letter and replace "
            "it with a token like `{{NAME}}`, then save that as your template."
        )
        st.write(
            "**Step 2:** Upload an Excel sheet where the column headers exactly match "
            "your tokens (without the curly braces) — e.g. columns named `NAME`, "
            "`DEPARTMENT`, `START_DATE`."
        )

        def _merge_one(template_bytes, replacements):
            doc = fitz.open(stream=template_bytes, filetype="pdf")
            for page in doc:
                blocks = [b for b in page.get_text("blocks") if b[4].strip()]
                text_dict = page.get_text("dict")
                dict_blocks = text_dict["blocks"]
                edits = []
                for b in blocks:
                    original = b[4]
                    new_text = original
                    changed = False
                    for token, val in replacements.items():
                        ph = "{{" + str(token) + "}}"
                        if ph in new_text:
                            new_text = new_text.replace(ph, str(val))
                            changed = True
                    if changed:
                        bbox = fitz.Rect(b[:4])
                        fontsize = 11
                        font_name = None
                        for db in dict_blocks:
                            for line in db.get("lines", []):
                                for span in line["spans"]:
                                    if fitz.Rect(span["bbox"]).intersects(bbox):
                                        fontsize = span["size"]
                                        font_name = span.get("font")
                        edits.append((bbox, new_text.rstrip("\n"), fontsize, font_name))
                page_rect = page.rect
                for bbox, new_text, fontsize, font_name in edits:
                    use_fontname = "helv"
                    if font_name:
                        try:
                            for xref, ext, subtype, basefont, refname, encoding in page.get_fonts(full=True):
                                if basefont == font_name or font_name in basefont:
                                    fn, fe, ft, fontbuffer = doc.extract_font(xref)
                                    if fontbuffer:
                                        alias = f"mergefont_{xref}"
                                        page.insert_font(fontname=alias, fontbuffer=fontbuffer)
                                        use_fontname = alias
                                    break
                        except Exception:
                            pass
                    try:
                        est_w = fitz.get_text_length(new_text, fontname=use_fontname, fontsize=fontsize)
                    except Exception:
                        est_w = fitz.get_text_length(new_text, fontname="helv", fontsize=fontsize)
                    # Padding is asymmetric: room above for ascenders, almost
                    # none below — avoids painting over a drawn underline rule
                    # that commonly sits just under a blank-field baseline.
                    box = fitz.Rect(
                        bbox.x0, bbox.y0 - fontsize * 0.3,
                        min(bbox.x0 + max(est_w, bbox.width) + 4, page_rect.width - 10),
                        bbox.y1 + fontsize * 0.05,
                    )
                    page.add_redact_annot(
                        box, text=new_text, fontname=use_fontname, fontsize=fontsize,
                        align=fitz.TEXT_ALIGN_LEFT, fill=(1, 1, 1), text_color=(0, 0, 0),
                    )
                page.apply_redactions()
            return doc.tobytes()

        template_pdf = st.file_uploader("Upload template PDF (with {{TOKENS}})", type=["pdf"], key="mm_template")
        excel_file = st.file_uploader("Upload Excel sheet (.xlsx)", type=["xlsx"], key="mm_excel")

        if template_pdf and excel_file:
            template_bytes = template_pdf.read()
            wb = openpyxl.load_workbook(io.BytesIO(excel_file.read()))
            ws = wb.active
            headers = [c.value for c in ws[1]]
            rows = list(ws.iter_rows(min_row=2, values_only=True))
            st.info(f"Found **{len(rows)}** rows and columns: {', '.join(str(h) for h in headers)}")

            folder_col = st.selectbox(
                "Which column should be used to sort output into folders? (e.g. department)",
                options=["(no folders — flat list)"] + [str(h) for h in headers],
                key="tok_folder_col",
            )
            name_col = st.selectbox(
                "Which column should be used to name each file? (e.g. name)",
                options=[str(h) for h in headers],
                key="tok_name_col",
            )

            if st.button(f"Generate all {len(rows)} personalized PDFs", key="tok_generate"):
                zip_buf = io.BytesIO()
                progress = st.progress(0, text="Generating...")
                with zipfile.ZipFile(zip_buf, "w") as zf:
                    for i, row in enumerate(rows):
                        rowdict = dict(zip(headers, row))
                        out_pdf = _merge_one(template_bytes, rowdict)
                        fname = str(rowdict.get(name_col, f"row_{i+1}")).replace(" ", "_").replace("/", "-")
                        if folder_col != "(no folders — flat list)":
                            folder = str(rowdict.get(folder_col, "General")).replace(" ", "_").replace("/", "-")
                            zf.writestr(f"{folder}/{fname}.pdf", out_pdf)
                        else:
                            zf.writestr(f"{fname}.pdf", out_pdf)
                        progress.progress((i + 1) / len(rows), text=f"Generated {i+1}/{len(rows)}")
                st.success(f"Done! Generated {len(rows)} personalized PDFs.")
                st.download_button(
                    "⬇️ Download all as ZIP",
                    data=zip_buf.getvalue(),
                    file_name="personalized_letters.zip",
                    mime="application/zip",
                    key="tok_download",
                )

    else:
        st.write(
            "**Best for:** any certificate/offer letter, in any layout, when you don't "
            "want to edit placeholder tokens into the file first. You upload the "
            "template as-is (PDF or image), pick an X/Y spot for each field **once**, "
            "and that exact same spot is used for all 2000 students — so the format "
            "never has to match a fixed pattern, and there's no risk of misaligned text."
        )

        st.write("#### 1. Get a demo Excel sheet")
        demo_headers = ["NAME", "DEPARTMENT", "COURSE_NAME", "DURATION"]
        demo_row = ["AJAY K", "B.Sc. AIML", "Cognitive Skills Enhancement - I", "June 2024 - November 2024"]
        _wb_demo = openpyxl.Workbook()
        _ws_demo = _wb_demo.active
        _ws_demo.title = "Students"
        _ws_demo.append(demo_headers)
        _ws_demo.append(demo_row)
        _buf_demo = io.BytesIO()
        _wb_demo.save(_buf_demo)
        st.download_button(
            "⬇️ Download demo Excel template (.xlsx)",
            data=_buf_demo.getvalue(),
            file_name="demo_student_data.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key="coord_demo_download",
        )
        st.caption("Edit this file: keep the same column headers, add one row per student. You can add more columns too.")

        st.write("#### 2. Upload your template and your real Excel sheet")
        coord_template = st.file_uploader(
            "Upload template (PDF or image — PNG/JPG)", type=["pdf", "png", "jpg", "jpeg"], key="coord_template"
        )
        coord_excel = st.file_uploader("Upload Excel sheet (.xlsx)", type=["xlsx"], key="coord_excel")

        if coord_template and coord_excel:
            # Get a base image of the template (render page 1 if PDF)
            if coord_template.type == "application/pdf" or coord_template.name.lower().endswith(".pdf"):
                _doc = fitz.open(stream=coord_template.read(), filetype="pdf")
                _pix = _doc[0].get_pixmap(matrix=fitz.Matrix(2, 2))
                base_image = Image.open(io.BytesIO(_pix.tobytes("png"))).convert("RGB")
            else:
                base_image = Image.open(coord_template).convert("RGB")

            W, H = base_image.size
            st.image(base_image, caption=f"Template preview ({W}×{H}px) — use these pixel coordinates below", width=500)

            wb = openpyxl.load_workbook(io.BytesIO(coord_excel.read()))
            ws = wb.active
            headers = [c.value for c in ws[1] if c.value is not None]
            rows = list(ws.iter_rows(min_row=2, values_only=True))
            st.info(f"Found **{len(rows)}** rows and columns: {', '.join(str(h) for h in headers)}")

            font_path_regular = None
            font_path_bold = None
            for candidate in [
                "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
            ]:
                if os.path.exists(candidate):
                    font_path_regular = candidate
                    break
            for candidate in [
                "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
                "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
            ]:
                if os.path.exists(candidate):
                    font_path_bold = candidate
                    break

            def _draw_positions_preview(sample_row):
                """Render the template with the CURRENT X/Y/size/align values
                stamped using one real row of data, so you can see exactly
                where text will land and fix it BEFORE running all 2000 —
                instead of guessing coordinates blind and finding out only
                after generating the whole batch."""
                rowdict = dict(zip(headers, sample_row))
                img_copy = base_image.copy()
                draw = ImageDraw.Draw(img_copy)
                for h, pos in positions.items():
                    text_val = str(rowdict.get(h, "") if rowdict.get(h) is not None else "")
                    chosen_path = font_path_bold if pos.get("bold") and font_path_bold else font_path_regular
                    try:
                        font = ImageFont.truetype(chosen_path, pos["size"]) if chosen_path else ImageFont.load_default()
                    except Exception:
                        font = ImageFont.load_default()
                    left, top, right, bottom = draw.textbbox((0, 0), text_val, font=font)
                    text_width = right - left
                    final_x = pos["x"] - (text_width / 2) if pos["align"] == "Center" else pos["x"]
                    # crosshair at the exact anchor point, so you can tell
                    # "position" from "where the text visually starts"
                    draw.line([(pos["x"] - 6, pos["y"]), (pos["x"] + 6, pos["y"])], fill=(255, 0, 0), width=1)
                    draw.line([(pos["x"], pos["y"] - 6), (pos["x"], pos["y"] + 6)], fill=(255, 0, 0), width=1)
                    draw.text((final_x, pos["y"]), text_val, fill=(0, 0, 0), font=font)
                return img_copy

            st.write("#### 3. Set the X/Y position for each field (same position used for every student)")
            positions = {}
            for idx, h in enumerate(headers):
                st.markdown(f"**Field: `{h}`**")
                cx, cy, cs, ca, cb = st.columns(5)
                with cx:
                    x_pos = st.number_input(f"X — {h}", min_value=0, max_value=W, value=int(W * 0.3), key=f"cx_{h}")
                with cy:
                    y_pos = st.number_input(f"Y — {h}", min_value=0, max_value=H, value=int(H * 0.5) + idx * 40, key=f"cy_{h}")
                with cs:
                    f_size = st.number_input(f"Font size — {h}", min_value=8, max_value=200, value=28, key=f"cs_{h}")
                with ca:
                    align_type = st.selectbox(f"Align — {h}", options=["Left", "Center"], key=f"ca_{h}")
                with cb:
                    bold_type = st.checkbox("Bold", key=f"cb_{h}")
                positions[h] = {"x": x_pos, "y": y_pos, "size": f_size, "align": align_type, "bold": bold_type}

            if rows:
                st.write("#### 👀 Live preview — check this BEFORE generating all rows")
                st.caption(
                    "Uses your first Excel row as a sample. The red crosshair marks the exact "
                    "X/Y anchor point for each field — adjust the numbers above and this updates."
                )
                st.image(_draw_positions_preview(rows[0]), caption="Preview with real data", width=500)

            folder_col2 = st.selectbox(
                "Sort output into folders by column? (optional)",
                options=["(no folders — flat list)"] + [str(h) for h in headers],
                key="coord_folder_col",
            )
            name_col2 = st.selectbox(
                "Use which column to name each file?", options=[str(h) for h in headers], key="coord_name_col"
            )


            if st.button(f"🚀 Generate all {len(rows)} personalized files", key="coord_generate"):
                zip_buf = io.BytesIO()
                progress = st.progress(0, text="Generating...")
                with zipfile.ZipFile(zip_buf, "w") as zf:
                    for i, row in enumerate(rows):
                        rowdict = dict(zip(headers, row))
                        img_copy = base_image.copy()
                        draw = ImageDraw.Draw(img_copy)
                        for h, pos in positions.items():
                            text_val = str(rowdict.get(h, "") if rowdict.get(h) is not None else "")
                            chosen_path = font_path_bold if pos.get("bold") and font_path_bold else font_path_regular
                            try:
                                font = ImageFont.truetype(chosen_path, pos["size"]) if chosen_path else ImageFont.load_default()
                            except Exception:
                                font = ImageFont.load_default()
                            left, top, right, bottom = draw.textbbox((0, 0), text_val, font=font)
                            text_width = right - left
                            final_x = pos["x"] - (text_width / 2) if pos["align"] == "Center" else pos["x"]
                            draw.text((final_x, pos["y"]), text_val, fill=(0, 0, 0), font=font)

                        pdf_out = io.BytesIO()
                        img_copy.save(pdf_out, format="PDF")
                        fname = str(rowdict.get(name_col2, f"file_{i+1}")).replace(" ", "_").replace("/", "-")
                        if folder_col2 != "(no folders — flat list)":
                            folder = str(rowdict.get(folder_col2, "General")).replace(" ", "_").replace("/", "-")
                            zf.writestr(f"{folder}/{fname}.pdf", pdf_out.getvalue())
                        else:
                            zf.writestr(f"{fname}.pdf", pdf_out.getvalue())
                        progress.progress((i + 1) / len(rows), text=f"Generated {i+1}/{len(rows)}")
                st.success(f"Done! Generated {len(rows)} personalized files.")
                st.download_button(
                    "⬇️ Download all as ZIP",
                    data=zip_buf.getvalue(),
                    file_name="personalized_letters.zip",
                    mime="application/zip",
                    key="coord_download",
                )

st.divider()
st.caption("HandyWriter · 100% free & open-source · Your files are processed locally, never uploaded anywhere.")
